"""
Generate a JEI-formatted Word manuscript (.docx) from the CAM project.

JEI formatting rules implemented:
- Arial 11 pt, 1.5 line spacing, 1-inch margins
- Continuous line numbers
- Section order: Title/Authors -> Summary -> Introduction -> Results ->
  Discussion -> Materials and Methods -> Acknowledgments -> References
- Figures placed at the end, each above its caption
- Headings are bold 11 pt (no special heading styles)
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "CAM_JEI_manuscript.docx")

FONT = "Arial"
SIZE = 11


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE)
    # ensure east-asian/complex also use Arial
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_line_numbers(doc):
    """Add continuous line numbers to the document body."""
    sectPr = doc.sections[0]._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sectPr.append(ln)


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(SIZE)
    return p


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(SIZE)
    return p


def title_line(doc, text, size, bold=True, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    return p


if __name__ == "__main__":
    from content import build  # noqa
    doc = Document()
    set_base_style(doc)
    add_line_numbers(doc)
    build(doc, DOCS, heading=heading, para=para, title_line=title_line)
    doc.save(OUT)
    print("WROTE", OUT)
